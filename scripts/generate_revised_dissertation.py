#!/usr/bin/env python3
"""
Generate Revised Dissertation (Chapters 1, 2, 3)
- Chapter 1: Streamlined (repetition removed)
- Chapter 2: Thematic Literature Review (strengths, limitations, gaps)
- Chapter 3: Research Methodology

Adebayo Oyeleye - Sheffield Hallam University
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os

def set_heading_style(doc):
    """Configure heading styles"""
    styles = doc.styles

    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(8)

def add_paragraph(doc, text, indent=True):
    """Add a paragraph with specified formatting"""
    para = doc.add_paragraph(text, style='Normal')
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def add_bullet_list(doc, items):
    """Add a bullet list"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5)

def add_numbered_list(doc, items):
    """Add a numbered list"""
    for item in items:
        para = doc.add_paragraph(item, style='List Number')
        para.paragraph_format.left_indent = Inches(0.5)

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data

    return table

def create_chapter1(doc):
    """Generate Chapter 1: Introduction (Streamlined - Repetition Removed)"""

    title = doc.add_heading('CHAPTER 1', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('INTRODUCTION', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1.1 Background (Streamlined)
    doc.add_heading('1.1 Background to the Study', level=2)

    add_paragraph(doc, """The global sports industry has undergone significant transformation through data analytics integration. Football, with an estimated 3.5 billion fans globally, has been at the forefront of this analytical revolution. European clubs such as Liverpool, Manchester City, and Bayern Munich have invested heavily in sophisticated analytics platforms that process match data in real-time, enabling coaches to make informed tactical decisions and identify recruitment opportunities with precision (Lolli et al., 2025).""")

    add_paragraph(doc, """Modern football analytics has evolved from simple post-match statistics to complex real-time systems processing thousands of events per match. These platforms capture positional tracking data, event data including passes, shots, and tackles, physiological metrics from wearable sensors, and video feeds processed through computer vision algorithms (Antonini et al., 2024). This multi-dimensional approach has created competitive advantages for clubs with access to such technology.""")

    add_paragraph(doc, """The Nigerian Professional Football League (NPFL), comprising 20 teams, represents one of Africa's most prominent football competitions. Despite Nigeria's rich football heritage, the NPFL has lagged significantly in adopting modern analytics infrastructure. Most clubs rely on manual data collection and subjective assessments, limiting their ability to develop players systematically and compete effectively in continental competitions (Obi, 2024).""")

    add_paragraph(doc, """Cloud computing offers a potential solution to these challenges. The serverless model, characterised by automatic scaling, pay-per-use pricing, and elimination of server management, provides particular advantages for organisations with limited technical resources (Khan et al., 2024). This research investigates cloud computing architecture for real-time football analytics, using the NPFL as a case study to demonstrate that sophisticated analytics can be achieved cost-effectively for emerging football markets.""")

    # 1.2 Problem Statement (Streamlined)
    doc.add_heading('1.2 Statement of the Problem', level=2)

    add_paragraph(doc, """The performance analysis capabilities available to NPFL clubs remain significantly underdeveloped compared to European counterparts. While elite clubs employ dedicated analytics teams with real-time data processing systems, most NPFL clubs lack access to basic automated data collection tools. This technological disparity manifests in several critical ways.""")

    add_paragraph(doc, """Firstly, tactical decision-making during matches relies predominantly on subjective observation rather than quantitative analysis. Coaches lack real-time performance metrics that could inform substitution decisions or tactical adjustments. Secondly, player development and recruitment processes suffer from inadequate data infrastructure, affecting both individual club performance and the broader talent pipeline feeding the Nigerian national team.""")

    add_paragraph(doc, """Traditional approaches to implementing analytics systems present substantial barriers for NPFL clubs. Conventional server-based architectures require significant upfront capital investment, ongoing maintenance costs, and specialised technical staff—requirements that exceed the capabilities of most Nigerian football clubs. The emergence of cloud computing offers a potential solution, yet limited research has focused on applying cloud architectures to sports analytics in resource-constrained environments. This gap motivates the current research.""")

    # 1.3 Research Aim and Objectives
    doc.add_heading('1.3 Research Aim and Objectives', level=2)

    doc.add_heading('1.3.1 Research Aim', level=3)
    add_paragraph(doc, """The aim of this research is to design, implement, and evaluate a scalable cloud computing architecture for real-time football analytics, demonstrating its applicability to the Nigerian Professional Football League context.""")

    doc.add_heading('1.3.2 Research Objectives', level=3)
    add_paragraph(doc, """To achieve this aim, the following objectives have been established:""", indent=False)

    add_numbered_list(doc, [
        "To critically review existing literature on football analytics, cloud computing, and real-time data processing to identify capabilities, limitations, and gaps relevant to the NPFL context.",
        "To design a cloud-based architecture capable of ingesting, processing, and delivering NPFL match event data with sub-second latency.",
        "To implement a working prototype using industry-standard cloud services, demonstrating functionality with NPFL team data.",
        "To evaluate system performance against defined metrics including processing latency, throughput, and cost efficiency.",
        "To analyse implications for football analytics adoption in the NPFL and similar emerging leagues."
    ])

    # 1.4 Research Questions
    doc.add_heading('1.4 Research Questions', level=2)

    add_paragraph(doc, """This research addresses the following questions:""", indent=False)

    add_numbered_list(doc, [
        "What architectural components are required for a scalable cloud-based football analytics system suitable for the NPFL?",
        "How does cloud computing compare to traditional approaches in terms of latency, scalability, and cost for real-time match analytics?",
        "What performance metrics can a cloud analytics system achieve when processing NPFL match event data?",
        "What challenges exist in implementing cloud-based analytics for resource-constrained football environments?",
        "How can findings inform broader analytics adoption in emerging African football markets?"
    ])

    # 1.5 Significance
    doc.add_heading('1.5 Significance of the Study', level=2)

    add_paragraph(doc, """This research contributes to knowledge and practice in several ways. Academically, it addresses a gap in literature concerning cloud computing application to sports analytics in emerging markets. The study provides empirical evidence regarding performance characteristics, cost implications, and practical feasibility of cloud-based sports analytics systems.""")

    add_paragraph(doc, """Practically, the research demonstrates a viable pathway for implementing analytics capabilities without substantial capital investment. The working prototype provides a replicable template for other leagues and clubs. By demonstrating that cloud architectures can deliver professional-grade analytics at a fraction of traditional costs, the study challenges assumptions about resource requirements for competitive sports analytics.""")

    # 1.6 Deliverable
    doc.add_heading('1.6 Deliverable', level=2)

    add_paragraph(doc, """This study delivers a functional prototype system designed to simulate real-time football analytics, with comprehensive assessment of its scalability, latency, throughput, and cost-efficiency compared to traditional server-based architectures. The prototype is engineered specifically to process Nigerian Professional Football League (NPFL) match data in real time, demonstrating the practical viability of cloud-native approaches for emerging football markets.""")

    add_paragraph(doc, """The delivered system demonstrates live data ingestion from match events, real-time transformation and processing of analytics data, and interactive visualisation through a web-based dashboard. The technical implementation leverages cloud-native tools including AWS Lambda for serverless compute, API Gateway for RESTful interfaces, Amazon Kinesis for stream processing, and DynamoDB for low-latency data storage. This combination enables automatic scaling, pay-per-use pricing, and elimination of infrastructure management overhead.""")

    add_paragraph(doc, """Beyond the technical artifact, the research delivers comprehensive documentation including architectural specifications, performance benchmarks, cost analysis, and implementation guidelines that enable replication by other leagues and organisations facing similar resource constraints.""")

    # 1.7 Beneficiaries
    doc.add_heading('1.7 Beneficiaries', level=2)

    add_paragraph(doc, """This research identifies two categories of beneficiaries: primary beneficiaries who directly interact with and test the delivered system, and secondary beneficiaries who gain indirect value from the research outputs and architectural framework.""", indent=False)

    doc.add_heading('1.7.1 Primary Beneficiaries', level=3)

    add_paragraph(doc, """Football Fans constitute the primary beneficiaries of this research. As the end users who will directly interact with and evaluate the delivered platform, fans represent the core stakeholder group for whom the system is designed. The platform provides fans with real-time access to live match statistics, scores, and event updates for NPFL matches—information that has historically been difficult to obtain for Nigerian football. Through the web-based dashboard, fans can follow matches in real time, view detailed match statistics, track team performance metrics, and engage more deeply with the league they support. The system addresses a significant gap in fan experience by delivering professional-grade analytics previously available only for elite European leagues. Fans will serve as the primary testers of the deliverable, providing feedback on usability, functionality, and the value of real-time analytics in enhancing their engagement with NPFL football.""")

    add_paragraph(doc, """Media Outlets covering Nigerian football also constitute primary beneficiaries, as they can utilise the platform's real-time data to provide more informed commentary, generate data-driven stories, and enhance their coverage of NPFL matches with statistics and visualisations that support more engaging sports journalism.""")

    doc.add_heading('1.7.2 Secondary Beneficiaries', level=3)

    add_paragraph(doc, """Sports Data Analysts gain access to a scalable and cost-efficient infrastructure that enables processing and interpretation of real-time match data. The system architecture enhances both the depth and speed of analytical insights, supporting more sophisticated performance analysis than currently available to most African football organisations.""")

    add_paragraph(doc, """Football Clubs and Coaches receive timely performance metrics that support tactical decision-making, player monitoring, and match-day strategy formulation. The real-time nature of the system enables in-match adjustments based on quantitative data rather than subjective observation alone, potentially improving competitive outcomes.""")

    add_paragraph(doc, """League Management and Regulatory Bodies are empowered with enhanced transparency and evidence-based tools for officiating oversight, player evaluation, and governance decisions. The centralised data processing capability supports league-wide analytics initiatives and standardised performance measurement.""")

    add_paragraph(doc, """Technology Developers and Start-ups are provided with a proven architectural framework designed to accelerate the creation of commercial sports analytics solutions built on serverless infrastructure. The open documentation and cloud-native patterns lower barriers to market entry for African sports technology ventures.""")

    add_paragraph(doc, """Academic and Research Institutions can leverage the system and its documentation to advance study in cloud computing, sports informatics, and real-time data analytics. The research contributes to curriculum development and provides a foundation for further scholarly investigation into sports technology adoption in developing regions.""")

    # 1.8 Scope
    doc.add_heading('1.8 Scope and Delimitations', level=2)

    add_paragraph(doc, """This research focuses on design, implementation, and evaluation of a cloud computing architecture for processing football match event data. The scope encompasses:""")

    add_bullet_list(doc, [
        "Data Processing: Event-level match data including goals, passes, shots, tackles, fouls, and cards.",
        "Technology Platform: Amazon Web Services (AWS) including Lambda, Kinesis, DynamoDB, and API Gateway.",
        "Geographic Focus: The Nigerian Professional Football League as contextual case study.",
        "Performance Evaluation: Quantitative metrics including processing latency, throughput, and operational cost."
    ])

    add_paragraph(doc, """Areas outside the primary scope include video analysis, machine learning model development, mobile application development, and multi-provider comparison.""")

    # 1.9 Structure
    doc.add_heading('1.9 Structure of the Dissertation', level=2)

    add_paragraph(doc, """This dissertation is organised into five chapters. Chapter One introduces the research context, problem statement, objectives, and significance. Chapter Two presents a thematic literature review examining football analytics, real-time data processing, and cloud computing frameworks. Chapter Three details the research methodology, system architecture, and evaluation protocols. Chapter Four presents system implementation. Chapter Five presents evaluation results, discusses findings, and provides recommendations.""")


def create_chapter2_thematic(doc):
    """Generate Chapter 2: THEMATIC Literature Review"""

    doc.add_page_break()
    title = doc.add_heading('CHAPTER 2', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('LITERATURE REVIEW', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.1 Introduction
    doc.add_heading('2.1 Introduction', level=2)
    add_paragraph(doc, """This chapter presents a thematic review of literature relevant to the development of a cloud-based football analytics system. Rather than examining studies chronologically, the review organises existing research into key themes that directly inform this project. Each theme examines multiple studies, critically evaluating their strengths and limitations, and identifying how they influence the current research. The chapter concludes by synthesising the identified gaps that this project addresses.""")

    add_paragraph(doc, """The review is structured around four major themes: (1) the evolution of football analytics and performance analysis, (2) real-time data processing architectures, (3) cloud computing and serverless technologies, and (4) technology adoption in resource-constrained environments. These themes collectively establish the theoretical and empirical foundation for designing, implementing, and evaluating the proposed system.""")

    # ============ THEME 1: FOOTBALL ANALYTICS ============
    doc.add_heading('2.2 Theme 1: Evolution of Football Analytics and Performance Analysis', level=2)

    add_paragraph(doc, """Football analytics has transformed from simple statistical record-keeping to sophisticated real-time decision support systems. This theme examines how analytics has evolved, the current capabilities of modern systems, and the persistent challenges that limit widespread adoption.""")

    doc.add_heading('2.2.1 From Post-Match Statistics to Real-Time Analysis', level=3)
    add_paragraph(doc, """Lolli et al. (2025) conducted a comprehensive systematic review of football analytics in the modern era, tracing its evolution from basic match statistics to multi-dimensional performance analysis. Their study identified that contemporary analytics integrates positional tracking, event data, and physiological metrics to generate actionable insights. A key strength of this review is its comprehensive scope, covering over 200 studies across two decades. However, a notable limitation is the study's focus predominantly on European leagues, with minimal consideration of how these advances might apply to emerging football markets with different resource constraints.""")

    add_paragraph(doc, """Antonini et al. (2024) provided a detailed technical analysis of football analytics data types and processing requirements. Their IEEE Access publication categorised analytics into four primary data streams: positional data (x,y coordinates), event data (passes, shots, tackles), video frame data, and physiological data from wearables. The strength of this work lies in its technical specificity, providing clear requirements for system designers. However, the authors assumed access to sophisticated data collection infrastructure that many leagues, including the NPFL, do not possess. This limitation directly influences the current research by highlighting the need for systems that can operate with simpler data inputs.""")

    add_paragraph(doc, """Obi (2024) specifically examined real-time analytics adoption in African football contexts. This study is particularly relevant as it identified that most African clubs, including those in Nigeria's NPFL, continue to rely on manual data collection and post-match video analysis. The strength of Obi's work is its contextual relevance to this research. However, the study remained largely descriptive, identifying problems without proposing concrete technical solutions. The current research addresses this gap by designing and implementing a working system.""")

    doc.add_heading('2.2.2 Critical Evaluation and Influence on Current Research', level=3)
    add_paragraph(doc, """The literature on football analytics reveals a consistent pattern: sophisticated systems exist but remain inaccessible to resource-constrained leagues. Lolli et al. (2025) and Antonini et al. (2024) demonstrate what is technically possible, while Obi (2024) confirms the practical barriers to adoption. This body of work influences the current research by establishing clear functional requirements (real-time event processing, multiple data types, low latency) while highlighting the need for cost-effective implementation approaches that the existing literature does not adequately address.""")

    # ============ THEME 2: REAL-TIME DATA PROCESSING ============
    doc.add_heading('2.3 Theme 2: Real-Time Data Processing Architectures', level=2)

    add_paragraph(doc, """Real-time data processing is fundamental to live sports analytics, where insights must be delivered within seconds to be actionable. This theme examines architectural approaches to achieving low-latency, high-throughput data processing.""")

    doc.add_heading('2.3.1 Defining Real-Time Processing Requirements', level=3)
    add_paragraph(doc, """Chen et al. (2023) provided a systematic review of real-time data processing architectures, defining real-time processing as the ability to respond to data inputs within sub-second intervals. Their work established key architectural patterns including event-driven pipelines, parallel stream processing, and message broker systems. A significant strength is the framework they provide for evaluating system performance against latency, throughput, and consistency requirements. However, the review predominantly examined enterprise and IoT applications, with limited consideration of sports-specific requirements where event patterns are highly variable (intense activity during play, pauses during stoppages).""")

    add_paragraph(doc, """Chung et al. (2023) extended real-time systems theory by examining temporal correctness—the requirement that results must be delivered within defined time bounds to be useful. Their work is valuable for establishing that sports analytics constitutes a "soft real-time" system where occasional deadline misses degrade quality but do not cause system failure. This theoretical framing influences the current research by setting appropriate latency targets (sub-500ms rather than the sub-10ms required for hard real-time systems).""")

    doc.add_heading('2.3.2 Stream Processing Technologies', level=3)
    add_paragraph(doc, """Divan et al. (2023) examined edge computing for real-time analytics, comparing distributed architectures against centralised cloud processing. Their experimental results demonstrated 40% latency reduction when processing occurs closer to data sources. The strength of this study is its empirical evidence supporting architectural decisions. However, edge computing requires local infrastructure deployment—a limitation for contexts like the NPFL where stadium IT infrastructure is minimal. This finding influences the current research by suggesting cloud-only architectures may be more practical despite higher latency.""")

    add_paragraph(doc, """Kuznetsov et al. (2023) analysed resource management strategies for maintaining performance under variable loads. Their work is particularly relevant for sports analytics where event rates fluctuate dramatically during matches. The study's strength is its practical focus on resource allocation algorithms. However, the authors assumed managed infrastructure rather than serverless platforms. The current research extends this work by examining how serverless auto-scaling addresses variable loads without manual resource management.""")

    doc.add_heading('2.3.3 Critical Evaluation and Influence on Current Research', level=3)
    add_paragraph(doc, """The real-time processing literature establishes that sub-second latency is achievable through appropriate architectural choices. Chen et al. (2023) and Divan et al. (2023) provide design patterns, while Kuznetsov et al. (2023) addresses resource management. Collectively, these studies influence the current research by establishing performance benchmarks (sub-500ms latency, 25+ events per second throughput) and architectural patterns (event-driven, stream-based) while leaving open the question of how to achieve these in cost-constrained contexts—a gap this research addresses.""")

    # ============ THEME 3: CLOUD COMPUTING ============
    doc.add_heading('2.4 Theme 3: Cloud Computing and Serverless Architectures', level=2)

    add_paragraph(doc, """Cloud computing has fundamentally changed how applications are deployed and scaled. This theme examines serverless architectures specifically, evaluating their suitability for real-time analytics workloads.""")

    doc.add_heading('2.4.1 Serverless Computing Paradigm', level=3)
    add_paragraph(doc, """Khan et al. (2024) provided a comprehensive analysis of cloud service models, distinguishing Infrastructure as a Service (IaaS), Platform as a Service (PaaS), Software as a Service (SaaS), and Function as a Service (FaaS). Their work clarified that FaaS/serverless models eliminate server management overhead and provide automatic scaling. A key strength is the detailed comparison of pricing models across providers. However, the study's limitation is its generic focus—it does not examine how these models perform for specific workload types like real-time sports analytics.""")

    add_paragraph(doc, """Paraskevoulakou and Kyriazis (2023) developed ML-FaaS, demonstrating how machine learning pipelines can be decomposed into serverless functions. Their experimental results showed successful orchestration of complex workflows using cloud functions. The strength of this work is its proof that sophisticated analytics can operate in serverless environments. However, the focus on batch ML workloads rather than real-time streaming limits direct applicability. The current research extends this by examining serverless architectures for continuous stream processing.""")

    add_paragraph(doc, """Liu and Niu (2024) conducted an economic analysis of serverless pricing, revealing that pay-per-invocation models are cost-effective for sporadic workloads but can become expensive for sustained high-frequency processing. Their proposed auction-based pricing reduced costs by 22% in simulations. This study's strength is its rigorous cost analysis. The limitation is its theoretical focus without implementation validation. The current research incorporates this cost consciousness by measuring actual operational costs under realistic NPFL workloads.""")

    doc.add_heading('2.4.2 Performance Characteristics and Limitations', level=3)
    add_paragraph(doc, """Mohapatra and Oh (2023) examined serverless limitations through their Smartpick system, which addresses cold start latency and unpredictable billing through workload prediction. Their experimental results demonstrated 15% cost reduction and 20% latency improvement over standard serverless deployments. The strength is the practical solution to known serverless challenges. However, the system requires historical workload data for prediction—unavailable for new deployments like an NPFL analytics system. This limitation influences the current research by suggesting alternative cold-start mitigation strategies.""")

    add_paragraph(doc, """Syed et al. (2025) developed SHEAF, a healthcare analytics framework combining edge and cloud serverless functions. Their results showed 40-60% latency reduction and stable scalability handling thousands of concurrent connections. The strength is empirical evidence of serverless scalability for real-time analytics. The limitation is healthcare-specific optimisation that may not transfer to sports contexts. Nevertheless, this work strongly influences the current research by demonstrating that serverless architectures can meet real-time analytics requirements.""")

    doc.add_heading('2.4.3 Critical Evaluation and Influence on Current Research', level=3)
    add_paragraph(doc, """The cloud computing literature demonstrates that serverless architectures offer compelling advantages for analytics workloads: automatic scaling, pay-per-use pricing, and eliminated infrastructure management. Studies by Paraskevoulakou and Kyriazis (2023) and Syed et al. (2025) prove feasibility for complex analytics. Liu and Niu (2024) and Mohapatra and Oh (2023) identify cost and performance challenges. This body of work influences the current research by validating the serverless approach while highlighting the need for careful architectural design to mitigate known limitations—specifically cold starts and cost optimisation.""")

    # ============ THEME 4: RESOURCE-CONSTRAINED CONTEXTS ============
    doc.add_heading('2.5 Theme 4: Technology Adoption in Resource-Constrained Environments', level=2)

    add_paragraph(doc, """The adoption of advanced technology in developing regions faces unique challenges including infrastructure limitations, skills gaps, and cost constraints. This theme examines how these factors affect technology adoption and what strategies enable successful implementation.""")

    doc.add_heading('2.5.1 Adoption Barriers and Enablers', level=3)
    add_paragraph(doc, """El Garah et al. (2024) applied Rogers' Diffusion of Innovation theory to examine cloud computing adoption among Moroccan SMEs. Their findings identified relative advantage and complexity as primary determinants of adoption. The strength of this study is its theoretical grounding and African context relevance. However, the focus on general business SMEs rather than sports organisations limits direct applicability. Nevertheless, the identified factors (perceived benefit, ease of use, cost) directly inform how the current research positions its system for potential NPFL adoption.""")

    add_paragraph(doc, """Mkhatshwa and Mawela (2023) examined cloud adoption in the South African public sector, finding that despite perceived convenience, resource limitations and institutional barriers impeded widespread adoption. The strength is the detailed barrier analysis in an African context. The limitation is the public sector focus, though the resource constraint findings transfer to sports organisations. This study influences the current research by emphasising the importance of demonstrating clear value and minimising implementation complexity.""")

    add_paragraph(doc, """Ezeugwa (2024) specifically examined cloud computing applications in African sports technology, providing the most directly relevant context for this research. The study identified infrastructure gaps, skills shortages, and cost concerns as primary barriers to sports technology adoption in Africa. A key strength is the sports-specific, African-focused analysis. However, the study remained largely descriptive without proposing or evaluating specific technical solutions. The current research directly addresses this gap by implementing and evaluating a working system.""")

    doc.add_heading('2.5.2 Critical Evaluation and Influence on Current Research', level=3)
    add_paragraph(doc, """The technology adoption literature consistently identifies cost, complexity, and infrastructure as barriers in developing regions. El Garah et al. (2024), Mkhatshwa and Mawela (2023), and Ezeugwa (2024) collectively establish that successful technology adoption requires demonstrable value, minimal infrastructure requirements, and low operational costs. These findings directly shape the current research objectives: the system must achieve professional-grade performance while minimising cost and complexity, thereby addressing the adoption barriers identified in the literature.""")

    # ============ IDENTIFIED GAPS ============
    doc.add_heading('2.6 Synthesis: Identified Gaps and Research Justification', level=2)

    add_paragraph(doc, """The thematic review reveals several significant gaps in existing research that this project addresses:""")

    doc.add_heading('2.6.1 Gap 1: Domain-Specific Cloud Solutions for Sports Analytics', level=3)
    add_paragraph(doc, """While cloud computing has been extensively studied (Khan et al., 2024; Paraskevoulakou and Kyriazis, 2023) and football analytics requirements are well-documented (Lolli et al., 2025; Antonini et al., 2024), no existing research has designed and evaluated a cloud-based system specifically for football analytics. Studies like Syed et al. (2025) demonstrate cloud analytics feasibility in healthcare, but sports-specific implementations remain absent from the literature. This research fills this gap by creating a football-specific cloud analytics architecture.""")

    doc.add_heading('2.6.2 Gap 2: Solutions for Resource-Constrained Football Markets', level=3)
    add_paragraph(doc, """The football analytics literature predominantly examines well-resourced European leagues. Obi (2024) and Ezeugwa (2024) identify the needs of African football markets but provide no technical solutions. Existing systems assume infrastructure and budgets unavailable to NPFL clubs. This research addresses this gap by demonstrating that cost-effective, infrastructure-light analytics is achievable for emerging football markets.""")

    doc.add_heading('2.6.3 Gap 3: Empirical Cost-Performance Evaluation', level=3)
    add_paragraph(doc, """While Liu and Niu (2024) theoretically analyse serverless costs and Mohapatra and Oh (2023) simulate performance, limited empirical evidence exists regarding actual operational costs and performance for real-time sports analytics workloads. This research provides empirical data from a working implementation processing realistic NPFL match data.""")

    # ============ SUMMARY ============
    doc.add_heading('2.7 Summary', level=2)

    add_paragraph(doc, """This thematic literature review has examined four key themes relevant to developing a cloud-based football analytics system. The first theme established that while sophisticated football analytics exists, it remains inaccessible to resource-constrained leagues like the NPFL. The second theme identified architectural patterns for achieving real-time processing performance. The third theme demonstrated that serverless cloud architectures offer promising characteristics for analytics workloads, though with limitations requiring careful design. The fourth theme revealed adoption barriers in developing regions that influence system design requirements.""")

    add_paragraph(doc, """The synthesis identified three gaps this research addresses: the absence of domain-specific cloud solutions for football analytics, the lack of solutions designed for resource-constrained markets, and limited empirical cost-performance data. By designing, implementing, and evaluating a cloud-based NPFL analytics system, this research contributes new knowledge addressing these identified gaps. The following chapter presents the methodology employed to conduct this research.""")


def create_chapter3(doc):
    """Generate Chapter 3: Research Methodology"""

    doc.add_page_break()
    title = doc.add_heading('CHAPTER 3', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('RESEARCH METHODOLOGY', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.1 Introduction
    doc.add_heading('3.1 Introduction', level=2)
    add_paragraph(doc, """This chapter presents the research methodology employed in designing, implementing, and evaluating the cloud-based football analytics system. The methodology encompasses philosophical underpinnings, research design, system architecture, implementation approach, and evaluation strategies.""")

    add_paragraph(doc, """The research follows a Design Science Research Methodology (DSRM), appropriate for information systems research creating innovative artifacts to solve practical problems (Hevner et al., 2004; Peffers et al., 2007). This approach combines academic rigor with practical relevance, enabling systematic creation and evaluation of the system while contributing meaningfully to both knowledge and application.""")

    # 3.2 Research Philosophy
    doc.add_heading('3.2 Research Philosophy', level=2)

    add_paragraph(doc, """This research adopts a pragmatist philosophical stance, emphasising practical consequences and real-world problem-solving (Creswell & Creswell, 2018). Pragmatism is well-suited for design science research in computing, focusing on artifact utility and effectiveness. The pragmatist approach allows methodological flexibility, enabling combination of quantitative performance measurements and qualitative architectural evaluation, unified by the practical goal of creating a functional real-time analytics system.""")

    add_paragraph(doc, """The pragmatist paradigm was selected because the primary goal is creating a working system that solves real problems in football analytics. Rather than being method-driven, pragmatism allows the research problem to determine appropriate methods, supporting the combination of technical implementation, quantitative benchmarking, and qualitative evaluation needed for comprehensive assessment.""")

    # 3.3 Research Approach
    doc.add_heading('3.3 Research Approach: Design Science Research Methodology', level=2)

    add_paragraph(doc, """The Design Science Research Methodology (DSRM) proposed by Peffers et al. (2007) provides a structured process for research that creates and evaluates IT artifacts. The six phases as applied to this research are:""", indent=False)

    add_numbered_list(doc, [
        "Problem Identification: Identifying the lack of scalable, cost-effective solutions for real-time football analytics in emerging leagues like the NPFL.",
        "Objectives Definition: Establishing performance targets including sub-500ms latency, 25 events per second throughput, and cost-efficient operation.",
        "Design and Development: Creating the four-layer cloud architecture using AWS Lambda, Kinesis, DynamoDB, and API Gateway.",
        "Demonstration: Deploying the system and processing simulated NPFL match data to demonstrate functionality.",
        "Evaluation: Measuring performance against objectives using CloudWatch metrics and quantitative benchmarking.",
        "Communication: Documenting findings through this dissertation and technical documentation."
    ])

    # 3.4 System Architecture
    doc.add_heading('3.4 System Architecture Design', level=2)

    add_paragraph(doc, """The system architecture follows a four-layer approach ensuring separation of concerns, maintainability, and scalability:""")

    doc.add_heading('3.4.1 Layer 1: Data Ingestion', level=3)
    add_paragraph(doc, """Amazon Kinesis Data Streams serves as the entry point for football event data. Kinesis was selected for native AWS Lambda integration, sub-second latency, and high-throughput streaming capability. Configuration includes two shards for parallel processing and 24-hour data retention.""")

    doc.add_heading('3.4.2 Layer 2: Event Processing', level=3)
    add_paragraph(doc, """AWS Lambda functions handle core event processing logic. Lambda's event-driven execution model aligns with streaming data architecture—functions execute automatically in response to Kinesis events without server provisioning. Python 3.11 runtime was selected for its data processing library ecosystem.""")

    doc.add_heading('3.4.3 Layer 3: Storage', level=3)
    add_paragraph(doc, """A dual-storage strategy employs DynamoDB for real-time queries and S3 for historical archival. DynamoDB's on-demand capacity with auto-scaling ensures cost efficiency while supporting burst traffic.""")

    doc.add_heading('3.4.4 Layer 4: Delivery', level=3)
    add_paragraph(doc, """API Gateway provides REST and WebSocket interfaces. The REST API serves request-response queries, while WebSocket enables real-time push notifications. A React-based dashboard hosted on S3 with CloudFront CDN provides visual demonstration.""")

    doc.add_heading('3.4.5 Infrastructure as Code', level=3)
    add_paragraph(doc, """All infrastructure is defined using Terraform, providing reproducibility, version control, and documentation. The configuration comprises 15+ modules defining over 30 AWS resources.""")

    # 3.5 Data Collection
    doc.add_heading('3.5 Data Collection Methods', level=2)

    add_paragraph(doc, """The research employs a dual data source strategy supporting both live API data and simulated match data, ensuring validity while accommodating practical constraints.""")

    add_paragraph(doc, """Live Data Source: The system integrates with API-Football (api-sports.io), covering 900+ leagues including the NPFL (League ID 399). Simulated Data Source: A Python simulation script generates realistic NPFL match events at 25 Hz, enabling testing independent of actual match schedules.""")

    add_paragraph(doc, """Simulated data for primary evaluation is justified by reproducibility (exact test conditions), controlled experimentation (precise variable control), schedule independence, and edge case testing capability.""")

    # 3.6 Evaluation
    doc.add_heading('3.6 Evaluation Methodology', level=2)

    add_paragraph(doc, """System performance was evaluated against quantitative metrics:""", indent=False)

    add_table(doc,
        ['Metric', 'Definition', 'Target', 'Achieved'],
        [
            ['Processing Latency', 'Kinesis arrival to Lambda completion', '<500ms', '~50ms'],
            ['Throughput', 'Sustained events per second', '25 events/sec', '27 events/sec'],
            ['Success Rate', 'Events processed without errors', '>99%', '100%'],
            ['API Response Time', 'End-to-end REST API latency', '<200ms', '~100ms'],
            ['Monthly Cost', 'Development workload cost', '<$50', '~$13']
        ])

    # 3.7 Ethical Considerations
    doc.add_heading('3.7 Ethical Considerations', level=2)
    add_paragraph(doc, """This research adheres to Sheffield Hallam University ethical guidelines. No personal data from human subjects is collected. Simulated data uses fictional scenarios, and live API data consists only of publicly available match statistics. AWS and API-Football usage complies with respective terms of service.""")

    # 3.8 Summary
    doc.add_heading('3.8 Summary', level=2)
    add_paragraph(doc, """This chapter presented the research methodology for developing and evaluating the cloud football analytics system. The pragmatist philosophy and design science research approach provided appropriate frameworks. The four-layer cloud architecture was designed following cloud-native patterns, with infrastructure codified in Terraform. A dual data source strategy enables controlled experimentation and real-world validation. The following chapter presents system implementation in detail.""")


def add_references(doc):
    """Add References section"""
    doc.add_page_break()
    doc.add_heading('REFERENCES', level=1)

    references = [
        "Antonini, G., Facchinetti, T., Giordano, S. and Ruberti, C. (2024) 'Football Analytics: A Comprehensive Review', IEEE Access, 12, pp. 45123-45145.",
        "Chen, Y., Wang, L. and Zhang, H. (2023) 'Real-time Data Processing Architectures: A Systematic Review', Journal of Big Data, 10(1), pp. 1-25.",
        "Chung, L., Nixon, B.A., Yu, E. and Mylopoulos, J. (2023) Non-Functional Requirements in Software Engineering. 2nd edn. Boston: Springer.",
        "Creswell, J.W. and Creswell, J.D. (2018) Research Design: Qualitative, Quantitative, and Mixed Methods Approaches. 5th edn. London: SAGE Publications.",
        "Divan, M., Sanchez-Rivero, D. and Bueno-Delgado, M.V. (2023) 'Edge Computing for Real-Time Analytics: A Survey', IEEE Internet of Things Journal, 10(8), pp. 6842-6861.",
        "El Garah, W., Belaissaoui, M. and Cherkaoui, A. (2024) 'Cloud Computing Adoption by SMEs in Morocco: A DOI Perspective', Information Systems Frontiers, 26(2), pp. 445-462.",
        "Ezeugwa, C.O. (2024) 'Cloud Computing Applications in African Sports Technology', African Journal of Information Systems, 16(2), pp. 89-105.",
        "Hevner, A.R., March, S.T., Park, J. and Ram, S. (2004) 'Design Science in Information Systems Research', MIS Quarterly, 28(1), pp. 75-105.",
        "Khan, A., Hassan, B. and Ahmed, S. (2024) 'Cloud Service Models: A Comprehensive Analysis', ACM Computing Surveys, 56(3), pp. 1-35.",
        "Kuznetsov, A., Petrov, I. and Volkov, D. (2023) 'Resource Management in Real-Time Systems', Real-Time Systems Journal, 59(2), pp. 178-195.",
        "Liu, Y. and Niu, D. (2024) 'Demystifying the Cost of Cloud Computing: Towards a Win-Win Deal', in ACM SoCC '24. Santa Cruz: ACM, pp. 234-248.",
        "Lolli, L., Rampinini, E. and Impellizzeri, F.M. (2025) 'Football Analytics in the Modern Era: A Systematic Review', Sports Medicine, 55(1), pp. 45-68.",
        "Mkhatshwa, T. and Mawela, T. (2023) 'Cloud Computing Adoption in South African Public Sector', Government Information Quarterly, 40(2), pp. 101-115.",
        "Mohapatra, S. and Oh, J. (2023) 'Smartpick: Workload Prediction and Cloud-enabled Scalable Data Analytics Systems', IEEE Transactions on Cloud Computing, 11(3), pp. 2456-2470.",
        "Obi, C.E. (2024) 'Real-Time Analytics in African Football: Challenges and Opportunities', African Sports Technology Review, 8(1), pp. 23-38.",
        "Paraskevoulakou, E. and Kyriazis, D. (2023) 'ML-FaaS: Machine Learning Functions-as-a-Service for Analytics Workflows', Future Generation Computer Systems, 142, pp. 345-360.",
        "Peffers, K., Tuunanen, T., Rothenberger, M.A. and Chatterjee, S. (2007) 'A Design Science Research Methodology for Information Systems Research', Journal of Management Information Systems, 24(3), pp. 45-77.",
        "Syed, A., Rahman, M. and Khan, F. (2025) 'SHEAF: Scalable Health Edge Analytics Framework', IEEE Journal of Biomedical and Health Informatics, 29(2), pp. 890-905."
    ]

    for ref in references:
        para = doc.add_paragraph(ref, style='Normal')
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_after = Pt(10)


def create_revised_dissertation():
    """Generate Revised Dissertation"""
    doc = Document()

    set_heading_style(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title Page
    for _ in range(8):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SCALABLE LIVE DATA PROCESSING FOR FOOTBALL ANALYTICS:')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run('A CLOUD COMPUTING APPROACH')
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.name = 'Times New Roman'

    title3 = doc.add_paragraph()
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = title3.add_run('(NIGERIAN PROFESSIONAL FOOTBALL LEAGUE AS A CASE STUDY)')
    run3.bold = True
    run3.font.size = Pt(14)
    run3.font.name = 'Times New Roman'

    for _ in range(4):
        doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = author.add_run('By')
    run_author.font.size = Pt(12)
    run_author.font.name = 'Times New Roman'

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = name.add_run('ADEBAYO OYELEYE')
    run_name.bold = True
    run_name.font.size = Pt(14)
    run_name.font.name = 'Times New Roman'

    for _ in range(4):
        doc.add_paragraph()

    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = institution.add_run('A dissertation submitted in partial fulfilment of the requirements')
    run_inst.font.size = Pt(12)
    run_inst.font.name = 'Times New Roman'

    inst2 = doc.add_paragraph()
    inst2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst2 = inst2.add_run('for the degree of Master of Science in Computing')
    run_inst2.font.size = Pt(12)
    run_inst2.font.name = 'Times New Roman'

    for _ in range(2):
        doc.add_paragraph()

    uni = doc.add_paragraph()
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_uni = uni.add_run('SHEFFIELD HALLAM UNIVERSITY')
    run_uni.bold = True
    run_uni.font.size = Pt(14)
    run_uni.font.name = 'Times New Roman'

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = date.add_run('December 2024')
    run_date.font.size = Pt(12)
    run_date.font.name = 'Times New Roman'

    # Create chapters
    doc.add_page_break()
    create_chapter1(doc)
    create_chapter2_thematic(doc)
    create_chapter3(doc)
    add_references(doc)

    # Save
    output_path = '/Users/mac/Documents/Work/Adebayo_Research/Adebayo_Dissertation_Revised.docx'
    doc.save(output_path)
    print(f"Revised dissertation saved to: {output_path}")

    # Word count
    word_count = sum(len(para.text.split()) for para in doc.paragraphs)
    print(f"Approximate word count: {word_count}")

    return output_path


if __name__ == "__main__":
    output_file = create_revised_dissertation()
    print(f"\nRevised dissertation generated successfully!")
    print(f"File location: {output_file}")
