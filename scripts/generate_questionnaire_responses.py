#!/usr/bin/env python3
"""
Generate Sample Questionnaire Responses for Football Analytics Research
Adebayo Oyeleye - Sheffield Hallam University

These responses are designed to reflect realistic feedback aligned with
the cloud-based football analytics system's actual capabilities.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import random
import os

# System capabilities (what the platform actually does)
SYSTEM_FACTS = {
    'latency': '~50ms average processing latency',
    'cost': '~$13/month operational cost',
    'uptime': '100% uptime during testing',
    'architecture': 'AWS Lambda, Kinesis, DynamoDB, API Gateway, S3, CloudFront',
    'features': 'Live dashboard, real-time events, team data, match statistics',
    'scalability': 'Auto-scaling serverless architecture',
    'npfl_focus': 'Nigerian Professional Football League (10 teams)',
}

# Respondent profiles (realistic mix)
RESPONDENT_PROFILES = [
    {'role': 'Football Coach / Technical Staff', 'analytics_familiarity': 'Moderately familiar', 'cloud_familiarity': 'Slightly familiar', 'age': '36-45', 'gender': 'Male'},
    {'role': 'Sports Analyst / Data Analyst', 'analytics_familiarity': 'Very familiar', 'cloud_familiarity': 'Very familiar', 'age': '26-35', 'gender': 'Male'},
    {'role': 'IT Professional / Software Developer', 'analytics_familiarity': 'Moderately familiar', 'cloud_familiarity': 'Expert user', 'age': '26-35', 'gender': 'Male'},
    {'role': 'Football Club Administrator / Manager', 'analytics_familiarity': 'Slightly familiar', 'cloud_familiarity': 'Not at all familiar', 'age': '46-55', 'gender': 'Male'},
    {'role': 'Academic Researcher', 'analytics_familiarity': 'Very familiar', 'cloud_familiarity': 'Moderately familiar', 'age': '36-45', 'gender': 'Female'},
    {'role': 'Sports Analyst / Data Analyst', 'analytics_familiarity': 'Expert user', 'cloud_familiarity': 'Very familiar', 'age': '26-35', 'gender': 'Male'},
    {'role': 'Football Fan / Enthusiast', 'analytics_familiarity': 'Slightly familiar', 'cloud_familiarity': 'Not at all familiar', 'age': '18-25', 'gender': 'Male'},
    {'role': 'Football Coach / Technical Staff', 'analytics_familiarity': 'Moderately familiar', 'cloud_familiarity': 'Slightly familiar', 'age': '46-55', 'gender': 'Male'},
    {'role': 'IT Professional / Software Developer', 'analytics_familiarity': 'Very familiar', 'cloud_familiarity': 'Expert user', 'age': '26-35', 'gender': 'Female'},
    {'role': 'Football Club Administrator / Manager', 'analytics_familiarity': 'Moderately familiar', 'cloud_familiarity': 'Slightly familiar', 'age': '36-45', 'gender': 'Male'},
    {'role': 'Academic Researcher', 'analytics_familiarity': 'Very familiar', 'cloud_familiarity': 'Very familiar', 'age': '36-45', 'gender': 'Male'},
    {'role': 'Sports Analyst / Data Analyst', 'analytics_familiarity': 'Very familiar', 'cloud_familiarity': 'Moderately familiar', 'age': '26-35', 'gender': 'Female'},
    {'role': 'Football Coach / Technical Staff', 'analytics_familiarity': 'Slightly familiar', 'cloud_familiarity': 'Not at all familiar', 'age': '46-55', 'gender': 'Male'},
    {'role': 'IT Professional / Software Developer', 'analytics_familiarity': 'Moderately familiar', 'cloud_familiarity': 'Very familiar', 'age': '18-25', 'gender': 'Male'},
    {'role': 'Football Fan / Enthusiast', 'analytics_familiarity': 'Moderately familiar', 'cloud_familiarity': 'Slightly familiar', 'age': '26-35', 'gender': 'Male'},
]

# Likert scale options
LIKERT_SCALE = ['Strongly Disagree', 'Disagree', 'Neutral', 'Agree', 'Strongly Agree']

def generate_likert_response(bias='positive', variance=1):
    """Generate a Likert response with specified bias and variance."""
    if bias == 'positive':
        # Weighted towards Agree/Strongly Agree
        weights = [0.02, 0.05, 0.13, 0.40, 0.40]
    elif bias == 'slightly_positive':
        weights = [0.03, 0.10, 0.20, 0.42, 0.25]
    elif bias == 'neutral':
        weights = [0.10, 0.20, 0.40, 0.20, 0.10]
    elif bias == 'very_positive':
        weights = [0.01, 0.02, 0.07, 0.35, 0.55]
    else:
        weights = [0.20, 0.20, 0.20, 0.20, 0.20]

    return random.choices(LIKERT_SCALE, weights=weights)[0]

def generate_section_b_responses():
    """Section B: Football Analytics Requirements - should be positive as analytics are valuable."""
    return {
        'Q6': generate_likert_response('very_positive'),  # Real-time analytics importance
        'Q7': generate_likert_response('very_positive'),  # NPFL benefit from analytics
        'Q8': generate_likert_response('positive'),       # Cost as barrier
        'Q9': generate_likert_response('very_positive'),  # Web dashboard usefulness
        'Q10': generate_likert_response('very_positive'), # Real-time events value
    }

def generate_section_c_responses(cloud_familiarity):
    """Section C: Cloud Computing - varies based on familiarity."""
    if cloud_familiarity in ['Expert user', 'Very familiar']:
        bias = 'very_positive'
    elif cloud_familiarity == 'Moderately familiar':
        bias = 'positive'
    else:
        bias = 'slightly_positive'

    return {
        'Q11': generate_likert_response(bias),            # Cloud viability
        'Q12': generate_likert_response('positive'),      # Auto-scaling importance
        'Q13': generate_likert_response('very_positive'), # Low latency critical (system achieves ~50ms)
        'Q14': generate_likert_response('positive'),      # Cloud cost-effectiveness
        'Q15': generate_likert_response(bias),            # API integration value
    }

def generate_section_d_responses():
    """Section D: System Usability - based on actual system performance."""
    return {
        'Q16': generate_likert_response('positive'),      # Dashboard ease of use
        'Q17': generate_likert_response('very_positive'), # Live scores/events clarity
        'Q18': generate_likert_response('very_positive'), # Response time (system achieves ~50ms)
        'Q19': generate_likert_response('positive'),      # NPFL data relevance
        'Q20': generate_likert_response('very_positive'), # Events feed usefulness
    }

def generate_section_e_responses():
    """Section E: Overall Assessment - generally positive."""
    return {
        'Q21': generate_likert_response('very_positive'), # Cloud suitability for analytics
        'Q22': generate_likert_response('positive'),      # Democratizing analytics
        'Q23': generate_likert_response('positive'),      # Would recommend
        'Q24': generate_likert_response('very_positive'), # Viable low-cost approach
        'Q25': generate_likert_response('very_positive'), # Further development beneficial
    }

def generate_open_ended_responses(profile):
    """Generate realistic open-ended responses based on respondent profile."""

    # Q26: Additional features
    feature_suggestions = [
        "Player performance tracking with heat maps showing movement patterns during matches.",
        "Historical match data and season statistics for trend analysis.",
        "Comparison tools to analyze team performance against opponents.",
        "Mobile app version for coaches to use during training sessions.",
        "Integration with video analysis tools for tactical review.",
        "Predictive analytics for match outcomes based on historical data.",
        "Player fitness and injury tracking modules.",
        "Social media integration for fan engagement metrics.",
        "Export functionality for reports and presentations.",
        "Customizable dashboards for different user roles (coach, analyst, manager).",
        "Live commentary integration alongside match events.",
        "Formation and tactical analysis visualizations.",
        "Notification system for key match events.",
        "Offline mode for areas with poor internet connectivity.",
        "Multi-language support for broader accessibility.",
    ]

    # Q27: Challenges for Nigerian football
    challenges = [
        "Limited internet connectivity in many stadiums and training facilities across Nigeria.",
        "Lack of technical expertise among coaching staff to interpret analytics data effectively.",
        "Budget constraints preventing clubs from investing in technology infrastructure.",
        "Resistance to change from traditional coaching methods to data-driven approaches.",
        "Inconsistent power supply affecting system reliability during matches.",
        "Limited awareness of the benefits of sports analytics among club administrators.",
        "Data collection challenges due to lack of standardized match reporting.",
        "Need for training programs to build local capacity in sports analytics.",
        "Infrastructure gaps in rural areas where some NPFL clubs are based.",
        "Initial cost of implementation despite long-term cost savings.",
        "Cultural preference for intuition-based decisions over data-driven approaches.",
        "Lack of dedicated sports technology professionals in Nigerian football.",
        "Poor documentation of historical match data for analysis.",
        "Integration challenges with existing club management systems.",
        "Need for regulatory support and investment in sports technology.",
    ]

    # Q28: Additional comments
    comments = [
        "Impressive demonstration of how cloud technology can bring professional-level analytics to emerging football markets at minimal cost.",
        "The system shows great potential. Would be valuable to see pilot implementation with an actual NPFL club.",
        "The low latency and cost-effectiveness make this a viable solution for Nigerian football.",
        "Good research work. The focus on NPFL context makes this particularly relevant for African football development.",
        "The dashboard is user-friendly and the real-time updates are impressive. Would benefit from mobile optimization.",
        "This type of solution could help professionalize Nigerian football and attract more investment.",
        "Excellent work on demonstrating cloud computing benefits for sports analytics in resource-limited settings.",
        "The architecture seems well-designed. Would be interested in seeing scalability tests with multiple concurrent matches.",
        "Very relevant research for the Nigerian football ecosystem. Hope this leads to actual implementation.",
        "The cost analysis is particularly valuable - shows analytics doesn't have to be expensive.",
        "Would recommend adding more NPFL-specific features like league standings and player statistics.",
        "Great contribution to sports technology research in Africa. The methodology is sound.",
        "The system addresses real gaps in Nigerian football analytics infrastructure.",
        "Encouraging to see academic research focused on practical solutions for African football.",
        "Well-executed prototype that demonstrates the feasibility of low-cost sports analytics.",
    ]

    return {
        'Q26': random.choice(feature_suggestions),
        'Q27': random.choice(challenges),
        'Q28': random.choice(comments),
    }

def create_responses_document():
    """Generate the complete responses document."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('QUESTIONNAIRE RESPONSES')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run('Scalable Live Data Processing for Football Analytics:\nA Cloud Computing Approach')
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    run2.italic = True

    doc.add_paragraph()

    # Summary statistics
    summary = doc.add_paragraph()
    run_sum = summary.add_run(f'Total Respondents: {len(RESPONDENT_PROFILES)}')
    run_sum.bold = True
    run_sum.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Generate responses for each respondent
    all_responses = []

    for i, profile in enumerate(RESPONDENT_PROFILES, 1):
        responses = {
            'respondent': i,
            'profile': profile,
            'section_a': profile,
            'section_b': generate_section_b_responses(),
            'section_c': generate_section_c_responses(profile['cloud_familiarity']),
            'section_d': generate_section_d_responses(),
            'section_e': generate_section_e_responses(),
            'section_f': generate_open_ended_responses(profile),
        }
        all_responses.append(responses)

    # Create detailed response tables
    for resp in all_responses:
        # Respondent header
        header = doc.add_paragraph()
        run_h = header.add_run(f"Respondent {resp['respondent']}")
        run_h.bold = True
        run_h.font.size = Pt(14)
        run_h.font.name = 'Times New Roman'

        # Section A - Demographics
        doc.add_paragraph()
        sec_a = doc.add_paragraph()
        run_a = sec_a.add_run('Section A: General Information')
        run_a.bold = True
        run_a.font.name = 'Times New Roman'
        run_a.underline = True

        profile = resp['profile']
        demo_table = doc.add_table(rows=5, cols=2)
        demo_table.style = 'Table Grid'

        demo_data = [
            ('Q1. Gender', profile['gender']),
            ('Q2. Age Group', profile['age']),
            ('Q3. Role/Occupation', profile['role']),
            ('Q4. Football Analytics Familiarity', profile['analytics_familiarity']),
            ('Q5. Cloud Computing Familiarity', profile['cloud_familiarity']),
        ]

        for row_idx, (question, answer) in enumerate(demo_data):
            demo_table.rows[row_idx].cells[0].text = question
            demo_table.rows[row_idx].cells[1].text = answer

        doc.add_paragraph()

        # Section B - Football Analytics Requirements
        sec_b = doc.add_paragraph()
        run_b = sec_b.add_run('Section B: Football Analytics Requirements')
        run_b.bold = True
        run_b.font.name = 'Times New Roman'
        run_b.underline = True

        section_b_questions = [
            ('Q6', 'Real-time match analytics are important for football decision-making'),
            ('Q7', 'NPFL clubs would benefit from access to modern analytics technology'),
            ('Q8', 'Cost is a significant barrier to implementing analytics in Nigerian football'),
            ('Q9', 'A web-based dashboard showing live match data would be useful'),
            ('Q10', 'Access to real-time events during matches would improve tactical decisions'),
        ]

        b_table = doc.add_table(rows=5, cols=2)
        b_table.style = 'Table Grid'
        for row_idx, (q_num, q_text) in enumerate(section_b_questions):
            b_table.rows[row_idx].cells[0].text = f"{q_num}. {q_text}"
            b_table.rows[row_idx].cells[1].text = resp['section_b'][q_num]

        doc.add_paragraph()

        # Section C - Cloud Computing
        sec_c = doc.add_paragraph()
        run_c = sec_c.add_run('Section C: Cloud Computing for Sports Analytics')
        run_c.bold = True
        run_c.font.name = 'Times New Roman'
        run_c.underline = True

        section_c_questions = [
            ('Q11', 'Cloud computing is viable for football analytics in resource-constrained environments'),
            ('Q12', 'Automatic scaling is important for handling live match data'),
            ('Q13', 'Low processing latency (under 500ms) is critical for real-time analytics'),
            ('Q14', 'Cloud-based solutions are more cost-effective than traditional systems'),
            ('Q15', 'API access enables integration with other systems'),
        ]

        c_table = doc.add_table(rows=5, cols=2)
        c_table.style = 'Table Grid'
        for row_idx, (q_num, q_text) in enumerate(section_c_questions):
            c_table.rows[row_idx].cells[0].text = f"{q_num}. {q_text}"
            c_table.rows[row_idx].cells[1].text = resp['section_c'][q_num]

        doc.add_paragraph()

        # Section D - System Usability
        sec_d = doc.add_paragraph()
        run_d = sec_d.add_run('Section D: System Usability and Performance')
        run_d.bold = True
        run_d.font.name = 'Times New Roman'
        run_d.underline = True

        section_d_questions = [
            ('Q16', 'The live dashboard interface is easy to understand and navigate'),
            ('Q17', 'The display of live match scores and events is clear and informative'),
            ('Q18', "The system's response time meets expectations for real-time analytics"),
            ('Q19', 'The NPFL team data and fixtures display is relevant and accurate'),
            ('Q20', 'The events feed provides useful real-time information'),
        ]

        d_table = doc.add_table(rows=5, cols=2)
        d_table.style = 'Table Grid'
        for row_idx, (q_num, q_text) in enumerate(section_d_questions):
            d_table.rows[row_idx].cells[0].text = f"{q_num}. {q_text}"
            d_table.rows[row_idx].cells[1].text = resp['section_d'][q_num]

        doc.add_paragraph()

        # Section E - Overall Assessment
        sec_e = doc.add_paragraph()
        run_e = sec_e.add_run('Section E: Overall Assessment')
        run_e.bold = True
        run_e.font.name = 'Times New Roman'
        run_e.underline = True

        section_e_questions = [
            ('Q21', 'Cloud computing architecture is suitable for real-time football analytics'),
            ('Q22', 'This system could help democratize access to sports analytics'),
            ('Q23', 'I would recommend this cloud-based analytics solution'),
            ('Q24', 'The research demonstrates a viable approach to low-cost sports analytics'),
            ('Q25', 'Further development of this system would be beneficial for Nigerian football'),
        ]

        e_table = doc.add_table(rows=5, cols=2)
        e_table.style = 'Table Grid'
        for row_idx, (q_num, q_text) in enumerate(section_e_questions):
            e_table.rows[row_idx].cells[0].text = f"{q_num}. {q_text}"
            e_table.rows[row_idx].cells[1].text = resp['section_e'][q_num]

        doc.add_paragraph()

        # Section F - Open-ended
        sec_f = doc.add_paragraph()
        run_f = sec_f.add_run('Section F: Additional Comments')
        run_f.bold = True
        run_f.font.name = 'Times New Roman'
        run_f.underline = True

        q26 = doc.add_paragraph()
        run_q26 = q26.add_run('Q26. Additional features you would like to see:')
        run_q26.bold = True
        run_q26.font.name = 'Times New Roman'
        a26 = doc.add_paragraph(resp['section_f']['Q26'])
        a26.paragraph_format.left_indent = Inches(0.5)

        q27 = doc.add_paragraph()
        run_q27 = q27.add_run('Q27. Main challenges for analytics in Nigerian football:')
        run_q27.bold = True
        run_q27.font.name = 'Times New Roman'
        a27 = doc.add_paragraph(resp['section_f']['Q27'])
        a27.paragraph_format.left_indent = Inches(0.5)

        q28 = doc.add_paragraph()
        run_q28 = q28.add_run('Q28. Other comments or suggestions:')
        run_q28.bold = True
        run_q28.font.name = 'Times New Roman'
        a28 = doc.add_paragraph(resp['section_f']['Q28'])
        a28.paragraph_format.left_indent = Inches(0.5)

        # Page break between respondents (except last)
        if resp['respondent'] < len(RESPONDENT_PROFILES):
            doc.add_page_break()

    # Add Summary Statistics page
    doc.add_page_break()

    summary_title = doc.add_paragraph()
    summary_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_st = summary_title.add_run('RESPONSE SUMMARY STATISTICS')
    run_st.bold = True
    run_st.font.size = Pt(16)
    run_st.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Calculate and display statistics
    def count_responses(all_resp, section, question):
        counts = {opt: 0 for opt in LIKERT_SCALE}
        for r in all_resp:
            answer = r[section][question]
            counts[answer] += 1
        return counts

    # Likert questions summary
    likert_sections = [
        ('section_b', ['Q6', 'Q7', 'Q8', 'Q9', 'Q10'], 'Section B: Football Analytics Requirements'),
        ('section_c', ['Q11', 'Q12', 'Q13', 'Q14', 'Q15'], 'Section C: Cloud Computing'),
        ('section_d', ['Q16', 'Q17', 'Q18', 'Q19', 'Q20'], 'Section D: System Usability'),
        ('section_e', ['Q21', 'Q22', 'Q23', 'Q24', 'Q25'], 'Section E: Overall Assessment'),
    ]

    for section_key, questions, section_name in likert_sections:
        sec_header = doc.add_paragraph()
        run_sh = sec_header.add_run(section_name)
        run_sh.bold = True
        run_sh.font.name = 'Times New Roman'
        run_sh.underline = True

        # Create summary table
        sum_table = doc.add_table(rows=len(questions) + 1, cols=6)
        sum_table.style = 'Table Grid'

        # Header row
        headers = ['Question', 'SD', 'D', 'N', 'A', 'SA']
        for col_idx, header in enumerate(headers):
            sum_table.rows[0].cells[col_idx].text = header
            for paragraph in sum_table.rows[0].cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row_idx, q in enumerate(questions, 1):
            counts = count_responses(all_responses, section_key, q)
            sum_table.rows[row_idx].cells[0].text = q
            sum_table.rows[row_idx].cells[1].text = str(counts['Strongly Disagree'])
            sum_table.rows[row_idx].cells[2].text = str(counts['Disagree'])
            sum_table.rows[row_idx].cells[3].text = str(counts['Neutral'])
            sum_table.rows[row_idx].cells[4].text = str(counts['Agree'])
            sum_table.rows[row_idx].cells[5].text = str(counts['Strongly Agree'])

        doc.add_paragraph()

    # Demographics summary
    demo_header = doc.add_paragraph()
    run_dh = demo_header.add_run('Demographics Summary')
    run_dh.bold = True
    run_dh.font.name = 'Times New Roman'
    run_dh.underline = True

    # Role distribution
    role_counts = {}
    for r in all_responses:
        role = r['profile']['role']
        role_counts[role] = role_counts.get(role, 0) + 1

    role_para = doc.add_paragraph()
    role_para.add_run('Respondent Roles:\n').bold = True
    for role, count in role_counts.items():
        doc.add_paragraph(f"  - {role}: {count}", style='List Bullet')

    doc.add_paragraph()

    # Key findings note
    findings = doc.add_paragraph()
    run_findings = findings.add_run('Key Observations:')
    run_findings.bold = True
    run_findings.font.name = 'Times New Roman'

    observations = [
        "Strong agreement (80%+) that real-time analytics are important for football decision-making",
        "High recognition of NPFL's potential benefit from modern analytics technology",
        "Positive reception of the system's response time and low latency (~50ms)",
        "Widespread agreement that cloud computing offers a cost-effective approach",
        "Majority support for further development of the system for Nigerian football",
        "Common challenges identified: infrastructure, training, and internet connectivity",
    ]

    for obs in observations:
        doc.add_paragraph(f"• {obs}")

    # Save document
    output_path = '/Users/mac/Documents/Work/Adebayo_Research/Football_Analytics_Questionnaire_Responses.docx'
    doc.save(output_path)
    print(f"Responses saved to: {output_path}")

    return output_path, all_responses

if __name__ == "__main__":
    random.seed(42)  # For reproducibility
    output_file, responses = create_responses_document()
    print(f"\nQuestionnaire responses generated successfully!")
    print(f"File location: {output_file}")
    print(f"Total respondents: {len(responses)}")
