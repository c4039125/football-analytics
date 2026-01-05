#!/usr/bin/env python3
"""
Generate Questionnaire for Football Analytics Research
Adebayo Oyeleye - Sheffield Hallam University
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os

def set_styles(doc):
    """Configure document styles"""
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

def add_checkbox_options(doc, options):
    """Add checkbox-style options"""
    for option in options:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run(f"☐ {option}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

def add_likert_scale(doc):
    """Add standard Likert scale options"""
    options = [
        "Strongly Disagree",
        "Disagree",
        "Neutral",
        "Agree",
        "Strongly Agree"
    ]
    add_checkbox_options(doc, options)

def add_question(doc, number, text, question_type='likert', options=None):
    """Add a question with appropriate response options"""
    para = doc.add_paragraph()
    run = para.add_run(f"{number}. {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

    if question_type == 'likert':
        add_likert_scale(doc)
    elif question_type == 'checkbox' and options:
        add_checkbox_options(doc, options)

    doc.add_paragraph()  # Add spacing

def create_questionnaire():
    """Generate the research questionnaire"""
    doc = Document()

    # Set up styles
    set_styles(doc)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('RESEARCH QUESTIONNAIRE')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run('Scalable Live Data Processing for Football Analytics:\nA Cloud Computing Approach')
    run2.font.size = Pt(14)
    run2.font.name = 'Times New Roman'
    run2.italic = True

    doc.add_paragraph()

    # Landing Page / Introduction
    intro_title = doc.add_paragraph()
    run_intro = intro_title.add_run('Introduction')
    run_intro.bold = True
    run_intro.font.size = Pt(12)
    run_intro.font.name = 'Times New Roman'

    intro_text = doc.add_paragraph()
    intro_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_text = intro_text.add_run(
        "You are being invited to participate in a research study titled 'Scalable Live Data Processing "
        "for Football Analytics: A Cloud Computing Approach'. This study is being conducted by Adebayo Oyeleye "
        "from the Department of Computing at Sheffield Hallam University."
    )
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(11)

    purpose = doc.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_purpose = purpose.add_run(
        "Purpose of this study: This research aims to evaluate the effectiveness of cloud computing "
        "architecture for real-time football analytics in the Nigerian Professional Football League (NPFL) context. "
        "Your responses will help assess the system's usability, performance, and potential impact on football "
        "analytics in emerging markets."
    )
    run_purpose.font.name = 'Times New Roman'
    run_purpose.font.size = Pt(11)

    what_asked = doc.add_paragraph()
    what_asked.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_asked = what_asked.add_run(
        "What you will be asked to do: Complete a short questionnaire about football analytics systems "
        "and cloud computing technology. This should take approximately 5-7 minutes."
    )
    run_asked.font.name = 'Times New Roman'
    run_asked.font.size = Pt(11)

    rights = doc.add_paragraph()
    rights.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_rights = rights.add_run(
        "Your rights: Your participation is entirely voluntary, and you can withdraw from the survey at any time "
        "by closing your web browser. You are free to skip any question you prefer not to answer."
    )
    run_rights.font.name = 'Times New Roman'
    run_rights.font.size = Pt(11)

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_conf = confidential.add_run(
        "Confidentiality: All responses will be collected anonymously and used solely for academic research purposes. "
        "No personally identifiable information will be collected or stored."
    )
    run_conf.font.name = 'Times New Roman'
    run_conf.font.size = Pt(11)

    doc.add_paragraph()

    # ==================== SECTION A ====================
    section_a = doc.add_paragraph()
    run_a = section_a.add_run('SECTION A: General Information')
    run_a.bold = True
    run_a.font.size = Pt(13)
    run_a.font.name = 'Times New Roman'
    run_a.underline = True

    doc.add_paragraph()

    add_question(doc, 1, "What is your gender?", 'checkbox', [
        "Male",
        "Female",
        "Other / Prefer not to say"
    ])

    add_question(doc, 2, "What is your age group?", 'checkbox', [
        "18-25",
        "26-35",
        "36-45",
        "46-55",
        "55+"
    ])

    add_question(doc, 3, "What is your primary role or occupation?", 'checkbox', [
        "Football Coach / Technical Staff",
        "Football Club Administrator / Manager",
        "Sports Analyst / Data Analyst",
        "IT Professional / Software Developer",
        "Academic Researcher",
        "Football Fan / Enthusiast",
        "Other (please specify): ________________"
    ])

    add_question(doc, 4, "How familiar are you with football analytics tools?", 'checkbox', [
        "Not at all familiar",
        "Slightly familiar",
        "Moderately familiar",
        "Very familiar",
        "Expert user"
    ])

    add_question(doc, 5, "How familiar are you with cloud computing technologies?", 'checkbox', [
        "Not at all familiar",
        "Slightly familiar",
        "Moderately familiar",
        "Very familiar",
        "Expert user"
    ])

    # ==================== SECTION B ====================
    doc.add_page_break()

    section_b = doc.add_paragraph()
    run_b = section_b.add_run('SECTION B: Football Analytics Requirements')
    run_b.bold = True
    run_b.font.size = Pt(13)
    run_b.font.name = 'Times New Roman'
    run_b.underline = True

    instruction_b = doc.add_paragraph()
    run_inst_b = instruction_b.add_run('Please indicate your level of agreement with the following statements:')
    run_inst_b.font.name = 'Times New Roman'
    run_inst_b.font.size = Pt(11)
    run_inst_b.italic = True

    doc.add_paragraph()

    add_question(doc, 6, "Real-time match analytics (live scores, events, statistics) are important for football decision-making.")

    add_question(doc, 7, "Nigerian Professional Football League (NPFL) clubs would benefit from access to modern analytics technology.")

    add_question(doc, 8, "Cost is a significant barrier to implementing analytics systems in Nigerian football.")

    add_question(doc, 9, "A web-based dashboard showing live match data would be useful for coaches and analysts.")

    add_question(doc, 10, "Access to real-time events (goals, cards, substitutions) during matches would improve tactical decisions.")

    # ==================== SECTION C ====================
    section_c = doc.add_paragraph()
    run_c = section_c.add_run('SECTION C: Cloud Computing for Sports Analytics')
    run_c.bold = True
    run_c.font.size = Pt(13)
    run_c.font.name = 'Times New Roman'
    run_c.underline = True

    instruction_c = doc.add_paragraph()
    run_inst_c = instruction_c.add_run('Please indicate your level of agreement with the following statements:')
    run_inst_c.font.name = 'Times New Roman'
    run_inst_c.font.size = Pt(11)
    run_inst_c.italic = True

    doc.add_paragraph()

    add_question(doc, 11, "Cloud computing (pay-as-you-go services) is a viable solution for football analytics in resource-constrained environments.")

    add_question(doc, 12, "Automatic scaling (system grows/shrinks based on demand) is important for handling live match data.")

    add_question(doc, 13, "Low processing latency (under 500ms response time) is critical for real-time sports analytics.")

    add_question(doc, 14, "Cloud-based solutions are more cost-effective than traditional server-based systems for small to medium organisations.")

    add_question(doc, 15, "The ability to access analytics via API (Application Programming Interface) enables integration with other systems.")

    # ==================== SECTION D ====================
    doc.add_page_break()

    section_d = doc.add_paragraph()
    run_d = section_d.add_run('SECTION D: System Usability and Performance')
    run_d.bold = True
    run_d.font.size = Pt(13)
    run_d.font.name = 'Times New Roman'
    run_d.underline = True

    instruction_d = doc.add_paragraph()
    run_inst_d = instruction_d.add_run(
        'The following questions relate to the cloud-based football analytics system developed in this research. '
        'If you have had the opportunity to view the system demo, please indicate your level of agreement:'
    )
    run_inst_d.font.name = 'Times New Roman'
    run_inst_d.font.size = Pt(11)
    run_inst_d.italic = True

    doc.add_paragraph()

    add_question(doc, 16, "The live dashboard interface is easy to understand and navigate.")

    add_question(doc, 17, "The display of live match scores and events is clear and informative.")

    add_question(doc, 18, "The system's response time (speed) meets expectations for real-time analytics.")

    add_question(doc, 19, "The NPFL team data and fixtures display is relevant and accurate.")

    add_question(doc, 20, "The events feed (goals, cards, shots) provides useful real-time information.")

    # ==================== SECTION E ====================
    section_e = doc.add_paragraph()
    run_e = section_e.add_run('SECTION E: Overall Assessment')
    run_e.bold = True
    run_e.font.size = Pt(13)
    run_e.font.name = 'Times New Roman'
    run_e.underline = True

    instruction_e = doc.add_paragraph()
    run_inst_e = instruction_e.add_run('Please indicate your level of agreement with the following statements:')
    run_inst_e.font.name = 'Times New Roman'
    run_inst_e.font.size = Pt(11)
    run_inst_e.italic = True

    doc.add_paragraph()

    add_question(doc, 21, "Cloud computing architecture is suitable for real-time football analytics applications.")

    add_question(doc, 22, "This type of system could help democratise access to sports analytics for emerging football markets.")

    add_question(doc, 23, "I would recommend this type of cloud-based analytics solution to football organisations.")

    add_question(doc, 24, "The research demonstrates a viable approach to implementing low-cost sports analytics.")

    add_question(doc, 25, "Further development of this system would be beneficial for Nigerian football.")

    # ==================== OPEN-ENDED QUESTIONS ====================
    doc.add_page_break()

    section_f = doc.add_paragraph()
    run_f = section_f.add_run('SECTION F: Additional Comments')
    run_f.bold = True
    run_f.font.size = Pt(13)
    run_f.font.name = 'Times New Roman'
    run_f.underline = True

    doc.add_paragraph()

    q26 = doc.add_paragraph()
    run_q26 = q26.add_run("26. What additional features would you like to see in a football analytics system?")
    run_q26.font.name = 'Times New Roman'
    run_q26.font.size = Pt(12)
    run_q26.bold = True

    # Add blank lines for response
    for _ in range(4):
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Inches(0.5)
        run_line = line.add_run("_" * 80)
        run_line.font.name = 'Times New Roman'
        run_line.font.size = Pt(11)

    doc.add_paragraph()

    q27 = doc.add_paragraph()
    run_q27 = q27.add_run("27. What do you consider the main challenges for implementing analytics in Nigerian football?")
    run_q27.font.name = 'Times New Roman'
    run_q27.font.size = Pt(12)
    run_q27.bold = True

    for _ in range(4):
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Inches(0.5)
        run_line = line.add_run("_" * 80)
        run_line.font.name = 'Times New Roman'
        run_line.font.size = Pt(11)

    doc.add_paragraph()

    q28 = doc.add_paragraph()
    run_q28 = q28.add_run("28. Any other comments or suggestions regarding the research or system?")
    run_q28.font.name = 'Times New Roman'
    run_q28.font.size = Pt(12)
    run_q28.bold = True

    for _ in range(4):
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Inches(0.5)
        run_line = line.add_run("_" * 80)
        run_line.font.name = 'Times New Roman'
        run_line.font.size = Pt(11)

    # ==================== FINAL PAGE ====================
    doc.add_paragraph()
    doc.add_paragraph()

    consent = doc.add_paragraph()
    consent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_consent = consent.add_run(
        "By completing and submitting this questionnaire, you confirm that you have read and understood "
        "the information provided above and consent to participate in this research study."
    )
    run_consent.font.name = 'Times New Roman'
    run_consent.font.size = Pt(11)
    run_consent.italic = True

    doc.add_paragraph()

    thanks = doc.add_paragraph()
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_thanks = thanks.add_run("Thank you for your participation!")
    run_thanks.font.name = 'Times New Roman'
    run_thanks.font.size = Pt(12)
    run_thanks.bold = True

    thanks2 = doc.add_paragraph()
    thanks2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_thanks2 = thanks2.add_run("Your responses will be kept confidential and used only for academic purposes.")
    run_thanks2.font.name = 'Times New Roman'
    run_thanks2.font.size = Pt(11)

    doc.add_paragraph()

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_contact = contact.add_run(
        "For questions about this research, please contact:\n"
        "Adebayo Oyeleye\n"
        "Email: Adebayo.I.Oyeleye@student.shu.ac.uk\n"
        "Sheffield Hallam University"
    )
    run_contact.font.name = 'Times New Roman'
    run_contact.font.size = Pt(10)

    # Save document
    output_path = '/Users/mac/Documents/Work/Adebayo_Research/Football_Analytics_Questionnaire.docx'
    doc.save(output_path)
    print(f"Questionnaire saved to: {output_path}")

    return output_path

if __name__ == "__main__":
    output_file = create_questionnaire()
    print(f"\nQuestionnaire generated successfully!")
    print(f"File location: {output_file}")
